from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document

from ..pipeline.models import TaskResult


def export_word(results: Iterable[TaskResult], output_path: Path, batch_name: str) -> Path:
    results_list = list(results)
    doc = Document()
    doc.add_heading("交付信息", level=1)
    doc.add_paragraph(f"客户/账号：{batch_name}")
    doc.add_paragraph(f"日期：{datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"总条数：{len(results_list)}")

    for idx, result in enumerate(results_list, start=1):
        doc.add_heading(f"视频 {idx}: {result.item.title}", level=2)
        if result.item.input_value:
            doc.add_paragraph(f"链接/来源：{result.item.input_value}")
        if result.summary:
            doc.add_paragraph(f"摘要：{result.summary}")
        if result.item.publish_timestamp:
            ts = result.item.publish_timestamp
            if ts > 10**11:
                ts = int(ts / 1000)
            publish_time = datetime.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M:%S")
            doc.add_paragraph(f"发布时间：{publish_time}")
        if result.item.duration_ms:
            duration_value = result.item.duration_ms
            total_seconds = int(duration_value if duration_value < 1000 else duration_value / 1000)
            minutes = total_seconds // 60
            seconds = total_seconds % 60
            doc.add_paragraph(f"时长：{minutes:02}:{seconds:02}")
        doc.add_paragraph("文案：")
        for para in result.transcript.text.split("\n"):
            doc.add_paragraph(para)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
